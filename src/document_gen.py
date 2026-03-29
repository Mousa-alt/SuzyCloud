"""Document generation utility — creates Word (.docx) and PDF files from text.

Used by Claude subprocess to generate downloadable documents on demand.
CLI usage: python -m src.document_gen "Title" < content.txt
Module usage: from src.document_gen import text_to_docx, text_to_pdf
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from src import config

logger = logging.getLogger(__name__)

OUTDIR = config.PROJECT_ROOT / "media" / "outgoing"


def _now() -> datetime:
    """Get current time, preferring Cairo timezone if available."""
    if hasattr(config, "cairo_now"):
        return config.cairo_now()
    return datetime.now()


def text_to_docx(text: str, title: str = "Document", filename: str | None = None) -> str | None:
    """Convert text (with WhatsApp-style formatting) to a Word document.

    Args:
        text: Content text. Supports *bold*, numbered lists, bullet points, URLs.
        title: Document title shown at the top.
        filename: Output filename (without path). Auto-generated from title if not given.

    Returns:
        Relative path (media/outgoing/filename.docx) on success, None on failure.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        # Title
        date_str = _now().strftime("%B %d, %Y")
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        title_p.add_run("\n")
        date_run = title_p.add_run(date_str)
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")

        # Parse content
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            # Section headers (fully *bold* lines)
            if re.match(r"^\*[^*]+\*$", stripped):
                header_text = stripped.strip("*")
                h = doc.add_paragraph()
                run = h.add_run(header_text)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                continue

            p = doc.add_paragraph()

            # Numbered items get slight indent
            if re.match(r"^\d+\.", stripped):
                p.paragraph_format.left_indent = Inches(0.2)

            # Bullet points
            if stripped.startswith(("- ", "• ")):
                p.style = "List Bullet"
                stripped = stripped[2:]

            # Process inline *bold* spans
            parts = re.split(r"(\*[^*]+\*)", stripped)
            for part in parts:
                if part.startswith("*") and part.endswith("*") and len(part) > 2:
                    run = p.add_run(part[1:-1])
                    run.bold = True
                    run.font.size = Pt(10.5)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10.5)

            # URLs get blue color
            for run in p.runs:
                if "http://" in run.text or "https://" in run.text:
                    run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)

        # Save
        OUTDIR.mkdir(parents=True, exist_ok=True)
        if not filename:
            safe_title = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "-")[:50] or "document"
            date_tag = _now().strftime("%Y-%m-%d")
            filename = f"{safe_title}-{date_tag}.docx"
        # Sanitize: strip path components to prevent traversal
        filename = Path(filename).name.replace("\n", "").replace("\r", "").strip()
        if not filename:
            filename = "document.docx"
        if not filename.lower().endswith(".docx"):
            filename += ".docx"

        outdir_resolved = OUTDIR.resolve()
        filepath = (outdir_resolved / filename).resolve()
        if filepath.parent != outdir_resolved:
            logger.warning(f"Rejected filename with path traversal: {filename}")
            return None
        doc.save(str(filepath))
        rel_path = f"media/outgoing/{filename}"
        logger.info(f"Generated document: {rel_path}")
        return rel_path

    except Exception:
        logger.exception("Failed to generate Word document")
        return None


def text_to_pdf(text: str, title: str = "Document", filename: str | None = None) -> str | None:
    """Convert text (with WhatsApp-style formatting) to a PDF document.

    Args:
        text: Content text. Supports *bold*, numbered lists, bullet points, URLs.
        title: Document title shown at the top.
        filename: Output filename (without path). Auto-generated from title if not given.

    Returns:
        Relative path (media/outgoing/filename.pdf) on success, None on failure.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        OUTDIR.mkdir(parents=True, exist_ok=True)
        if not filename:
            safe_title = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "-")[:50] or "document"
            date_tag = _now().strftime("%Y-%m-%d")
            filename = f"{safe_title}-{date_tag}.pdf"
        filename = Path(filename).name.replace("\n", "").replace("\r", "").strip()
        if not filename:
            filename = "document.pdf"
        if not filename.lower().endswith(".pdf"):
            filename += ".pdf"

        outdir_resolved = OUTDIR.resolve()
        filepath = (outdir_resolved / filename).resolve()
        if filepath.parent != outdir_resolved:
            logger.warning(f"Rejected filename with path traversal: {filename}")
            return None

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DocTitle", parent=styles["Title"], fontSize=20,
            textColor=colors.HexColor("#1a1a2e"), spaceAfter=4,
            alignment=TA_CENTER,
        )
        date_style = ParagraphStyle(
            "DocDate", parent=styles["Normal"], fontSize=11,
            textColor=colors.HexColor("#666666"), alignment=TA_CENTER,
            spaceAfter=16,
        )
        heading_style = ParagraphStyle(
            "SectionHead", parent=styles["Heading2"], fontSize=14,
            textColor=colors.HexColor("#2c3e50"), spaceBefore=12, spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "Body", parent=styles["Normal"], fontSize=10.5,
            leading=14, spaceAfter=4,
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=body_style, leftIndent=20,
            bulletIndent=8, spaceAfter=3,
        )
        numbered_style = ParagraphStyle(
            "Numbered", parent=body_style, leftIndent=20,
            spaceAfter=3,
        )
        link_color = "#2980B9"

        doc = SimpleDocTemplate(
            str(filepath), pagesize=A4,
            topMargin=2 * cm, bottomMargin=2 * cm,
            leftMargin=2.2 * cm, rightMargin=2.2 * cm,
        )

        elements = []
        date_str = _now().strftime("%B %d, %Y")
        elements.append(Paragraph(xml_escape(title), title_style))
        elements.append(Paragraph(xml_escape(date_str), date_style))
        elements.append(Spacer(1, 8))

        def _inline_bold(line: str) -> str:
            """Convert *bold* spans to <b>bold</b> and URLs to blue links."""
            # Escape XML-special chars first, then apply markup
            parts = re.split(r"(\*[^*]+\*)", line)
            result = []
            for part in parts:
                if part.startswith("*") and part.endswith("*") and len(part) > 2:
                    result.append(f"<b>{xml_escape(part[1:-1])}</b>")
                else:
                    result.append(xml_escape(part))
            escaped = "".join(result)
            escaped = re.sub(
                r"(https?://\S+)",
                rf'<font color="{link_color}">\1</font>',
                escaped,
            )
            return escaped

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                elements.append(Spacer(1, 6))
                continue

            # Section headers (fully *bold* lines)
            if re.match(r"^\*[^*]+\*$", stripped):
                header_text = stripped.strip("*")
                elements.append(Paragraph(xml_escape(header_text), heading_style))
                continue

            # Bullet points
            if stripped.startswith(("- ", "• ")):
                content_text = _inline_bold(stripped[2:])
                elements.append(Paragraph(f"\u2022 {content_text}", bullet_style))
                continue

            # Numbered items
            m = re.match(r"^(\d+\.)\s*(.*)$", stripped)
            if m:
                num_prefix = xml_escape(m.group(1))
                content_text = _inline_bold(m.group(2))
                elements.append(Paragraph(f"{num_prefix} {content_text}", numbered_style))
                continue

            elements.append(Paragraph(_inline_bold(stripped), body_style))

        doc.build(elements)
        rel_path = f"media/outgoing/{filename}"
        logger.info(f"Generated PDF: {rel_path}")
        return rel_path

    except Exception:
        logger.exception("Failed to generate PDF document")
        return None


if __name__ == "__main__":
    """CLI: python -m src.document_gen "Document Title" [--pdf] < content.txt

    Or with explicit filename:
        python -m src.document_gen "Title" --filename report.docx < content.txt
    """
    import sys

    if len(sys.argv) < 2:
        print("Usage: python -m src.document_gen 'Title' [--filename name.ext] [--pdf] < content.txt")
        sys.exit(1)

    doc_title = sys.argv[1]
    doc_filename = None
    use_pdf = "--pdf" in sys.argv
    if "--filename" in sys.argv:
        idx = sys.argv.index("--filename")
        if idx + 1 >= len(sys.argv):
            print("Error: --filename requires a value", file=sys.stderr)
            sys.exit(1)
        doc_filename = sys.argv[idx + 1]

    content = sys.stdin.read()
    if not content.strip():
        print("Error: no content on stdin")
        sys.exit(1)

    if use_pdf:
        result = text_to_pdf(content, title=doc_title, filename=doc_filename)
    else:
        result = text_to_docx(content, title=doc_title, filename=doc_filename)
    if result:
        print(result)
    else:
        print("Error: document generation failed", file=sys.stderr)
        sys.exit(1)
